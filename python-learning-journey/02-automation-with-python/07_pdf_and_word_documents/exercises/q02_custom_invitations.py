"""
Question:
Given a list of guest names, generate a separate .docx invitation file
for each guest, with their name inserted into a template message.

T
"""

import os
from docx import Document


def create_invitations(guest_names, output_folder="invitations"):
    os.makedirs(output_folder, exist_ok=True)

    for name in guest_names:
        doc = Document()
        doc.add_heading("You're Invited!", level=1)
        doc.add_paragraph(f"Dear {name},")
        doc.add_paragraph(
            "You're invited to celebrate the launch of my Python automation journey! "
            "Join us for snacks, code, and good company."
        )
        doc.add_paragraph("Looking forward to seeing you there.")

        filename = os.path.join(output_folder, f"invitation_{name.replace(' ', '_')}.docx")
        doc.save(filename)
        print(f"Created: {filename}")

    print(f"\nDone. {len(guest_names)} invitation(s) created in '{output_folder}'.")


if __name__ == "__main__":
    guests = ["Rahul Sharma", "Priya Verma", "Amit Kumar"]
    create_invitations(guests)